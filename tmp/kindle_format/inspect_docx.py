from __future__ import annotations

import json
import sys
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def val(value):
    if value is None:
        return None
    try:
        return int(value)
    except TypeError:
        return str(value)


def has_page_break(paragraph):
    return bool(paragraph._p.xpath(".//w:br[@w:type='page']"))


def run_font_info(run):
    rpr = run._r.rPr
    fonts = {}
    if rpr is not None and rpr.rFonts is not None:
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            fonts[key] = rpr.rFonts.get(qn(f"w:{key}"))
    return {
        "text": run.text[:60],
        "bold": run.bold,
        "italic": run.italic,
        "size_pt": run.font.size.pt if run.font.size else None,
        "name": run.font.name,
        "fonts": fonts,
    }


def para_info(paragraph):
    fmt = paragraph.paragraph_format
    return {
        "style": paragraph.style.name if paragraph.style else None,
        "text": paragraph.text[:160],
        "full_len": len(paragraph.text),
        "alignment": val(paragraph.alignment),
        "left_indent_pt": fmt.left_indent.pt if fmt.left_indent else None,
        "first_line_indent_pt": fmt.first_line_indent.pt if fmt.first_line_indent else None,
        "space_before_pt": fmt.space_before.pt if fmt.space_before else None,
        "space_after_pt": fmt.space_after.pt if fmt.space_after else None,
        "line_spacing": fmt.line_spacing,
        "page_break": has_page_break(paragraph),
        "runs": [run_font_info(r) for r in paragraph.runs[:4]],
    }


def section_info(document):
    out = []
    for section in document.sections:
        out.append(
            {
                "page_width_pt": section.page_width.pt,
                "page_height_pt": section.page_height.pt,
                "top_margin_pt": section.top_margin.pt,
                "bottom_margin_pt": section.bottom_margin.pt,
                "left_margin_pt": section.left_margin.pt,
                "right_margin_pt": section.right_margin.pt,
                "header_distance_pt": section.header_distance.pt,
                "footer_distance_pt": section.footer_distance.pt,
            }
        )
    return out


def style_snapshot(document, used_styles):
    result = {}
    for name in sorted(used_styles):
        try:
            style = document.styles[name]
        except Exception:
            continue
        pf = style.paragraph_format
        font = style.font
        result[name] = {
            "type": str(style.type),
            "base": style.base_style.name if style.base_style else None,
            "font_name": font.name,
            "font_size_pt": font.size.pt if font.size else None,
            "bold": font.bold,
            "italic": font.italic,
            "alignment": val(pf.alignment),
            "left_indent_pt": pf.left_indent.pt if pf.left_indent else None,
            "first_line_indent_pt": pf.first_line_indent.pt if pf.first_line_indent else None,
            "space_before_pt": pf.space_before.pt if pf.space_before else None,
            "space_after_pt": pf.space_after.pt if pf.space_after else None,
            "line_spacing": pf.line_spacing,
            "keep_with_next": pf.keep_with_next,
            "page_break_before": pf.page_break_before,
        }
    return result


def zip_counts(path):
    with zipfile.ZipFile(path) as zf:
        names = zf.namelist()
    return {
        "images": len([n for n in names if n.startswith("word/media/")]),
        "headers": len([n for n in names if n.startswith("word/header") and n.endswith(".xml")]),
        "footers": len([n for n in names if n.startswith("word/footer") and n.endswith(".xml")]),
        "footnotes": "word/footnotes.xml" in names,
        "endnotes": "word/endnotes.xml" in names,
        "comments": "word/comments.xml" in names,
        "custom_props": "docProps/custom.xml" in names,
    }


def main(path_str):
    path = Path(path_str)
    doc = Document(path)
    styles = Counter(p.style.name if p.style else "<none>" for p in doc.paragraphs)
    style_lengths = defaultdict(list)
    page_breaks = []
    for i, p in enumerate(doc.paragraphs):
        style_lengths[p.style.name if p.style else "<none>"].append(len(p.text))
        if has_page_break(p):
            page_breaks.append(i)

    samples = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() or has_page_break(p):
            item = para_info(p)
            item["index"] = i
            samples.append(item)
        if len(samples) >= 80:
            break

    tail_samples = []
    for i in range(max(0, len(doc.paragraphs) - 30), len(doc.paragraphs)):
        p = doc.paragraphs[i]
        item = para_info(p)
        item["index"] = i
        tail_samples.append(item)

    output = {
        "path": str(path),
        "paragraph_count": len(doc.paragraphs),
        "table_count": len(doc.tables),
        "section_count": len(doc.sections),
        "sections": section_info(doc),
        "zip_counts": zip_counts(path),
        "style_counts": styles.most_common(),
        "style_avg_lengths": {
            k: round(sum(v) / len(v), 2) for k, v in sorted(style_lengths.items())
        },
        "page_break_paragraph_indexes": page_breaks[:200],
        "used_style_snapshot": style_snapshot(doc, set(styles.keys())),
        "samples": samples,
        "tail_samples": tail_samples,
    }
    print(json.dumps(output, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main(sys.argv[1])

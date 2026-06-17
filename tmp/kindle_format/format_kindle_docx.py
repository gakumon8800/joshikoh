from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt


FONT_NAME = "BIZ UDPゴシック"
SERIES_TITLE = "女子高生でもわかる不動産"
VOLUME_TITLE = "第5巻：先生、賃貸トラブルは契約書だけでは終わりません"
SUBTITLE = "―感情ではなく、契約・手続・証拠で動くための賃貸判例8選―"


def set_rfonts(element, font_name: str) -> None:
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)


def set_style_font(style, font_name: str, size_pt: float | None = None, bold=None) -> None:
    style.font.name = font_name
    set_rfonts(style._element, font_name)
    if size_pt is not None:
        style.font.size = Pt(size_pt)
    if bold is not None:
        style.font.bold = bold


def set_run_font(run, size_pt: float | None = None, bold=None) -> None:
    run.font.name = FONT_NAME
    set_rfonts(run._element, FONT_NAME)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def clear_paragraph(paragraph) -> None:
    paragraph._p.clear_content()


def replace_paragraph_text(paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    paragraph.add_run(text)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for key, value in values.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int, col_widths: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = table._tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in col_widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)


def format_styles(doc: Document) -> None:
    styles = doc.styles

    normal = styles["Normal"]
    set_style_font(normal, FONT_NAME, 10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    for name, size, before, after, align in [
        ("Title", 20, 0, 0, WD_ALIGN_PARAGRAPH.CENTER),
        ("Subtitle", 14, 0, 0, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 1", 16, 14, 4, WD_ALIGN_PARAGRAPH.CENTER),
        ("Heading 2", 14, 8, 4, WD_ALIGN_PARAGRAPH.LEFT),
        ("Heading 3", 12, 8, 4, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        style = styles[name]
        set_style_font(style, FONT_NAME, size, True if name.startswith("Heading") else None)
        pf = style.paragraph_format
        pf.alignment = align
        pf.first_line_indent = Pt(0)
        pf.left_indent = Pt(0)
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)
        pf.line_spacing = 1.0
        if name.startswith("Heading"):
            pf.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        set_style_font(style, FONT_NAME, 10.5)
        pf = style.paragraph_format
        pf.left_indent = Pt(19.85)
        pf.first_line_indent = Pt(-7.1)
        pf.space_before = Pt(0)
        pf.space_after = Pt(3)
        pf.line_spacing = 1.0


def format_sections(doc: Document) -> None:
    for section in doc.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(35)
        section.bottom_margin = Mm(30)
        section.left_margin = Mm(30)
        section.right_margin = Mm(30)
        section.header_distance = Mm(15)
        section.footer_distance = Mm(17.5)
        section.different_first_page_header_footer = False

        for story in (section.header, section.first_page_header, section.even_page_header):
            for paragraph in story.paragraphs:
                clear_paragraph(paragraph)

        for story in (section.footer, section.first_page_footer, section.even_page_footer):
            for paragraph in story.paragraphs:
                clear_paragraph(paragraph)
            for table in story.tables:
                table._element.getparent().remove(table._element)


def format_title_page(doc: Document) -> None:
    paragraphs = doc.paragraphs
    if len(paragraphs) < 6:
        return

    replace_paragraph_text(paragraphs[0], SERIES_TITLE)
    paragraphs[0].style = doc.styles["Normal"]
    paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraphs[0].paragraph_format.first_line_indent = Pt(0)
    paragraphs[0].paragraph_format.space_after = Pt(0)
    set_run_font(paragraphs[0].runs[0], 20, False)

    replace_paragraph_text(paragraphs[2], VOLUME_TITLE)
    paragraphs[2].style = doc.styles["Normal"]
    paragraphs[2].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraphs[2].paragraph_format.first_line_indent = Pt(0)
    set_run_font(paragraphs[2].runs[0], 14, False)

    replace_paragraph_text(paragraphs[3], SUBTITLE)
    paragraphs[3].style = doc.styles["Normal"]
    paragraphs[3].alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraphs[3].paragraph_format.first_line_indent = Pt(0)
    set_run_font(paragraphs[3].runs[0], 10.5, True)

    for idx in (1, 4):
        paragraphs[idx].style = doc.styles["Normal"]
        paragraphs[idx].alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraphs[idx].paragraph_format.first_line_indent = Pt(0)


def collect_toc_indexes(paragraphs) -> set[int]:
    toc_indexes: set[int] = set()
    in_toc = False
    for idx, paragraph in enumerate(paragraphs):
        style = paragraph.style.name if paragraph.style else ""
        text = paragraph.text.strip()
        if style == "Heading 1":
            in_toc = text == "目次"
            continue
        if in_toc:
            if paragraph._p.xpath(".//w:br[@w:type='page']"):
                in_toc = False
                continue
            toc_indexes.add(idx)
    return toc_indexes


def format_paragraphs(doc: Document) -> None:
    paragraphs = doc.paragraphs
    toc_indexes = collect_toc_indexes(paragraphs)
    for idx, paragraph in enumerate(paragraphs):
        style_name = paragraph.style.name if paragraph.style else ""
        pf = paragraph.paragraph_format

        if idx in (0, 2, 3):
            continue

        if style_name in ("Title", "Subtitle"):
            paragraph.style = doc.styles["Normal"]
            style_name = "Normal"

        if style_name == "Heading 1":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(0)
            pf.space_before = Pt(14)
            pf.space_after = Pt(4)
            pf.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, 16, True)
            continue

        if style_name == "Heading 2":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(0)
            pf.space_before = Pt(8)
            pf.space_after = Pt(4)
            pf.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, 14, True)
            continue

        if style_name == "Heading 3":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.first_line_indent = Pt(0)
            pf.left_indent = Pt(0)
            pf.space_before = Pt(8)
            pf.space_after = Pt(4)
            pf.keep_with_next = True
            for run in paragraph.runs:
                set_run_font(run, 12, True)
            continue

        if style_name in ("List Bullet", "List Number"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pf.space_before = Pt(0)
            pf.space_after = Pt(3)
            pf.line_spacing = 1.0
            for run in paragraph.runs:
                set_run_font(run, 10.5)
            continue

        if idx <= 5 or idx in toc_indexes:
            pf.first_line_indent = Pt(0)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx > 5 else paragraph.alignment
        else:
            pf.first_line_indent = Pt(10.5)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing = 1.0
        for run in paragraph.runs:
            set_run_font(run, 10.5)


def format_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        col_widths = [1900, 1900, 4700]
        set_table_width(table, sum(col_widths), col_widths)

        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_width(cell, col_widths[min(col_idx, len(col_widths) - 1)])
                set_cell_margins(cell)
                if row_idx == 0:
                    set_cell_shading(cell, "EDEDED")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER if col_idx < 2 else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    paragraph.paragraph_format.first_line_indent = Pt(0)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        set_run_font(run, 9.5, True if row_idx == 0 else None)


def has_page_break(paragraph) -> bool:
    return any(
        br.get(qn("w:type")) == "page" for br in paragraph._p.iter(qn("w:br"))
    )


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def compact_page_breaks(doc: Document) -> None:
    paragraphs = list(doc.paragraphs)
    remove_ids: set[int] = set()

    for idx, paragraph in enumerate(paragraphs):
        if not has_page_break(paragraph) or paragraph.text.strip():
            continue

        previous_idx = idx - 1
        while previous_idx >= 0:
            previous = paragraphs[previous_idx]
            if has_page_break(previous):
                break
            if previous.text.strip():
                previous.add_run().add_break(WD_BREAK.PAGE)
                remove_ids.add(id(paragraph))
                for blank_idx in range(previous_idx + 1, idx):
                    blank = paragraphs[blank_idx]
                    if not blank.text.strip() and not has_page_break(blank):
                        remove_ids.add(id(blank))
                break
            previous_idx -= 1

    for paragraph in paragraphs:
        if id(paragraph) in remove_ids:
            remove_paragraph(paragraph)


def set_core_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = VOLUME_TITLE
    props.subject = SUBTITLE.strip("―")
    props.keywords = "Kindle, 不動産, 賃貸トラブル, 判例"
    props.comments = ""


def main(input_path: str, output_path: str) -> None:
    doc = Document(input_path)
    format_styles(doc)
    format_sections(doc)
    format_title_page(doc)
    format_paragraphs(doc)
    format_tables(doc)
    compact_page_breaks(doc)
    set_core_properties(doc)

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("usage: format_kindle_docx.py input.docx output.docx")
    main(sys.argv[1], sys.argv[2])
